from __future__ import annotations

import argparse
import ctypes
import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import webbrowser
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
import pypandoc

sys.path.insert(0, str(Path(__file__).resolve().parent))

from git_source import (  # noqa: E402  (local module; import needs the path above)
    FILESYSTEM_SOURCE_TYPE,
    FILESYSTEM_SOURCE_TYPES,
    GIT_SOURCE_TYPE,
    GitSourceError,
    GitSourceSpec,
    parse_git_source,
    verify_generated_document,
    write_generated_git_source,
)
from source_snapshot import (  # noqa: E402  (local module; import needs the path above)
    SourceIsolationError,
    assert_source_isolation,
    build_snapshot,
    render_table,
)

#: The only document this change is authorised to re-source, and the exact Git
#: coordinates it must resolve through. Every other entry must keep resolving
#: from the redesign checkout.
DEPLOY_STATE_KEY = "deploy-state"
GIT_SOURCED_KEYS = frozenset({DEPLOY_STATE_KEY})

#: Documents whose resolved file is selected dynamically rather than named by
#: the manifest. The active-handoff entry tracks the newest handoff, so its
#: filename legitimately rotates; its root and Drive identity are still pinned.
ROTATING_KEYS = frozenset({"latest-handoff"})
REQUIRED_GIT_SOURCES = {
    DEPLOY_STATE_KEY: (
        r"C:\Users\izzyh\Documents\Terraforming Mars",
        "fix/live-compare-data-remove-declared-style",
    )
}


APP_DIR = Path(__file__).resolve().parent
# The redesign checkout. This is the only working-tree root, and it is fixed:
# no environment variable may redirect it, so a same-named file in another
# checkout can never become a planning-pack source.
ROOT = Path(r"C:\Users\izzyh\Documents\Terraforming Mars Redesign")
LOCAL_DRIVE_FOLDER = Path(r"G:\My Drive\TM Project Planning Pack")
FOLDER_NAME = "TM Project Planning Pack"
SCOPES = ["https://www.googleapis.com/auth/drive.file"]
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
GOOGLE_FOLDER_MIME = "application/vnd.google-apps.folder"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
APP_MARKER = "tm-planning-pack-v1"

CREDENTIALS_FILE = APP_DIR / "credentials.json"
TOKEN_FILE = APP_DIR / "token.json"
STATE_FILE = APP_DIR / "drive-map.json"
SUMMARY_FILE = APP_DIR / "last-run-summary.json"
REFERENCE_DOCX = APP_DIR / "google-docs-reference.docx"
SANITIZER = APP_DIR / "google_docs_title_sanitize.py"
LOCK_FILE = APP_DIR / "update.lock"
LOG_DIR = APP_DIR / "logs"
GENERATED_DIR = APP_DIR / "generated"
MASTER_CONTEXT_SOURCE = GENERATED_DIR / "tm-project-master-context.md"

# Where the catalog JSON itself is read from. `--source-manifest` may point this
# at a manifest committed in an isolated worktree for a single run.
#
# This selects the manifest FILE only. It does not re-root anything: every
# filesystem entry still resolves against ROOT above, phase files still come
# from ROOT, and the master-context and handoff sources still come from ROOT.
# An overriding manifest therefore cannot redirect a document to another
# checkout, and cannot change any non-DEPLOY-STATE source configuration.
DEFAULT_SOURCE_MANIFEST = ROOT / "docs/redesign/CLAUDE-PROJECT-SOURCES.json"
SOURCE_MANIFEST = DEFAULT_SOURCE_MANIFEST


def set_source_manifest(path: Path | None) -> Path:
    """Select the manifest file for this run. Returns the resolved path."""

    global SOURCE_MANIFEST
    if path is None:
        SOURCE_MANIFEST = DEFAULT_SOURCE_MANIFEST
        return SOURCE_MANIFEST
    candidate = path.expanduser().resolve()
    if not candidate.is_file():
        raise RuntimeError(f"--source-manifest is not a file: {candidate}")
    SOURCE_MANIFEST = candidate
    logging.warning(
        "Reading the source catalog from an overriding manifest: %s "
        "(document roots are unaffected; filesystem sources still resolve under %s)",
        candidate,
        ROOT,
    )
    return SOURCE_MANIFEST


@dataclass(frozen=True)
class Source:
    key: str
    title: str
    path: Path
    #: Set when the document is resolved from Git rather than from a working
    #: tree. ``path`` then points at the generated artifact, never at a source
    #: checkout, and the spec is re-verified against Git before publication.
    git_spec: GitSourceSpec | None = None


class UpdateAlreadyRunning(RuntimeError):
    pass


def read_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def repo_relative(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError as exc:
        raise RuntimeError(f"Master-context source is outside the redesign repository: {path}") from exc


def current_phase_path(state_text: str) -> Path:
    section = re.search(
        r"(?ms)^## Current substep\s*$\n(.*?)(?=^##\s|\Z)", state_text
    )
    if not section:
        raise RuntimeError("REDESIGN_STATE.md has no 'Current substep' section.")
    phase_match = re.search(r"\bPhase\s+(\d+)\s*,\s*Step\b", section.group(1))
    if not phase_match:
        raise RuntimeError("Could not determine the current phase from REDESIGN_STATE.md.")
    phase_number = int(phase_match.group(1))
    matches = sorted((ROOT / "docs/redesign/phases").glob(f"{phase_number:02d}-*.md"))
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one current phase {phase_number:02d} Markdown file; "
            f"found {len(matches)}."
        )
    return matches[0]


def resolve_handoff_path(relative_path: str) -> Path:
    normalized = relative_path.replace("\\", "/")
    if not normalized.startswith("docs/agent-handoffs/"):
        raise RuntimeError(f"Active handoff is outside docs/agent-handoffs: {relative_path}")
    candidate = (ROOT / Path(*normalized.split("/"))).resolve()
    handoff_root = (ROOT / "docs/agent-handoffs").resolve()
    try:
        candidate.relative_to(handoff_root)
    except ValueError as exc:
        raise RuntimeError(f"Active handoff escapes docs/agent-handoffs: {relative_path}") from exc
    if not candidate.is_file():
        raise RuntimeError(f"Active handoff does not exist: {relative_path}")
    return candidate


def active_handoff_paths(state_text: str) -> list[Path]:
    lines = state_text.splitlines()
    try:
        start = next(
            index for index, line in enumerate(lines) if line.strip() == "## Latest handoff"
        )
    except StopIteration as exc:
        raise RuntimeError("REDESIGN_STATE.md has no 'Latest handoff' section.") from exc

    paths: list[Path] = []
    list_started = False
    for line in lines[start + 1 :]:
        stripped = line.strip()
        if stripped.startswith("## "):
            break
        if not stripped:
            if list_started:
                break
            continue
        if stripped.startswith("- "):
            match = re.match(
                r"^-\s+`?(docs/agent-handoffs/[^`\s]+\.md)`?(?:\s|$)", stripped
            )
            if not match:
                if list_started:
                    raise RuntimeError(
                        "The active handoff block contains an unparseable list item: " + stripped
                    )
                continue
            list_started = True
            paths.append(resolve_handoff_path(match.group(1)))
        elif list_started:
            # Wrapped handoff descriptions belong to the current list item.
            continue

    if not paths:
        raise RuntimeError("The first handoff group under 'Latest handoff' is empty.")
    if len({path.resolve() for path in paths}) != len(paths):
        raise RuntimeError("The active handoff block contains duplicate paths.")
    return paths


def newest_handoff_path() -> Path:
    handoffs = list((ROOT / "docs/agent-handoffs").glob("*.md"))
    if not handoffs:
        raise RuntimeError("No Markdown handoffs exist in docs/agent-handoffs.")
    return max(handoffs, key=lambda path: (path.stat().st_mtime_ns, path.name.casefold()))


def demote_markdown_headings(markdown: str, levels: int = 2) -> str:
    output: list[str] = []
    fence: str | None = None
    for line in markdown.splitlines():
        fence_match = re.match(r"^\s*(`{3,}|~{3,})", line)
        if fence_match:
            marker = fence_match.group(1)[0]
            if fence is None:
                fence = marker
            elif fence == marker:
                fence = None
            output.append(line)
            continue
        if fence is None:
            heading = re.match(r"^(#{1,6})(\s+.*)$", line)
            if heading:
                level = min(6, len(heading.group(1)) + levels)
                line = "#" * level + heading.group(2)
        output.append(line)
    return "\n".join(output).rstrip() + "\n"


def master_context_sources() -> tuple[list[Path], Path, list[Path], Path]:
    guide = ROOT / "docs/redesign/CLAUDE-PROJECT-CONTEXT.md"
    current_status = ROOT / "docs/CURRENT_STATUS.md"
    authority_index = ROOT / "docs/AUTHORITATIVE_DOCUMENTS.md"
    state = ROOT / "docs/REDESIGN_STATE.md"
    required = (guide, current_status, authority_index, state)
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise RuntimeError("Missing master-context source files:\n" + "\n".join(missing))
    state_text = read_markdown(state)
    phase = current_phase_path(state_text)
    active_handoffs = active_handoff_paths(state_text)
    newest = newest_handoff_path()
    handoffs = list(active_handoffs)
    if newest.resolve() not in {path.resolve() for path in handoffs}:
        handoffs.append(newest)
    ordered = [guide, current_status, authority_index, state, phase, *handoffs]
    return ordered, phase, active_handoffs, newest


def build_master_context_source() -> Path:
    ordered, phase, active_handoffs, newest = master_context_sources()
    digest = hashlib.sha256()
    for path in ordered:
        relative = repo_relative(path)
        content = path.read_bytes()
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        digest.update(content)
        digest.update(b"\0")
    fingerprint = digest.hexdigest()

    lines = [
        "# TM PROJECT MASTER CONTEXT",
        "",
        "> Permanent generated orientation page for the Claude Project. Add this native "
        "Google Doc once; the updater refreshes the same Drive file ID.",
        "",
        f"> Source snapshot fingerprint: `{fingerprint}`",
        "",
        "## Read this first",
        "",
        "1. Read the concise current status and authoritative document index embedded "
        "below before planning or modifying code.",
        "2. Then read the full current state, active phase, and active handoff set.",
        "3. Follow the authority order in the canonical root instructions and master rules.",
        "4. Treat this generated page as navigation and aggregation only. It never overrides "
        "a canonical source and never authorizes broader scope, production work, the next "
        "substep, a deployment, a push, or a migration.",
        "5. When a statement here conflicts with an attached canonical source, follow the "
        "canonical source and flag the generated page as stale.",
        "",
        "## Included source manifest",
        "",
        "- Context contract: `docs/redesign/CLAUDE-PROJECT-CONTEXT.md`",
        "- Concise current status: `docs/CURRENT_STATUS.md`",
        "- Authority and evidence index: `docs/AUTHORITATIVE_DOCUMENTS.md`",
        "- Current state: `docs/REDESIGN_STATE.md`",
        f"- Current phase: `{repo_relative(phase)}`",
        f"- Declared active handoffs: {len(active_handoffs)}",
    ]
    for handoff in active_handoffs:
        lines.append(f"  - `{repo_relative(handoff)}`")
    if newest.resolve() not in {path.resolve() for path in active_handoffs}:
        lines.extend(
            [
                "- Newest repository handoff, included as a freshness backstop:",
                f"  - `{repo_relative(newest)}`",
            ]
        )
    else:
        lines.append("- Newest repository handoff is already in the declared active set.")
    lines.extend(["", "---", ""])

    labels = {
        ordered[0]: "Canonical context contract",
        ordered[1]: "Concise current status",
        ordered[2]: "Authority and evidence index",
        ordered[3]: "Canonical current state",
        ordered[4]: "Canonical current phase",
    }
    active_set = {path.resolve() for path in active_handoffs}
    for path in ordered:
        if path.resolve() in active_set:
            label = "Declared active handoff"
        elif path == newest and path not in labels:
            label = "Newest repository handoff freshness backstop"
        else:
            label = labels.get(path, "Canonical source")
        lines.extend(
            [
                f"## {label}: `{repo_relative(path)}`",
                "",
                demote_markdown_headings(read_markdown(path)).rstrip(),
                "",
                "---",
                "",
            ]
        )

    content = "\n".join(lines).rstrip() + "\n"
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    temp = MASTER_CONTEXT_SOURCE.with_suffix(".md.tmp")
    temp.write_text(content, encoding="utf-8", newline="\n")
    os.replace(temp, MASTER_CONTEXT_SOURCE)
    return MASTER_CONTEXT_SOURCE


def load_source_manifest() -> dict[str, Any]:
    if not SOURCE_MANIFEST.is_file():
        raise RuntimeError(f"Planning-pack source manifest is missing: {SOURCE_MANIFEST}")
    try:
        manifest = json.loads(SOURCE_MANIFEST.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise RuntimeError(f"Planning-pack source manifest is invalid: {SOURCE_MANIFEST}") from exc
    if manifest.get("version") != 1:
        raise RuntimeError("Planning-pack source manifest must use version 1.")
    if not isinstance(manifest.get("documents"), list):
        raise RuntimeError("Planning-pack source manifest has no documents array.")
    if not isinstance(manifest.get("phaseDocuments"), dict):
        raise RuntimeError("Planning-pack source manifest has no phaseDocuments object.")
    if not isinstance(manifest.get("dynamicDocuments"), list):
        raise RuntimeError("Planning-pack source manifest has no dynamicDocuments array.")
    return manifest


def resolve_manifest_path(root_name: Any, relative_value: Any, label: str) -> Path:
    # ``redesign`` is the only working-tree root. A document owned by another
    # repository lineage must be declared as a Git source; there is deliberately
    # no root that points at another checkout's working tree, because such a
    # copy is stale whenever that lineage is committed to from elsewhere.
    roots = {"redesign": ROOT}
    if root_name not in roots:
        raise RuntimeError(f"{label} has an unknown root: {root_name!r}")
    if not isinstance(relative_value, str) or not relative_value.strip():
        raise RuntimeError(f"{label} has no relative path.")
    relative = Path(relative_value)
    if relative.is_absolute() or ".." in relative.parts:
        raise RuntimeError(f"{label} path must stay inside its declared root: {relative_value}")
    root = roots[root_name].resolve()
    candidate = (root / relative).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise RuntimeError(f"{label} path escapes its declared root: {relative_value}") from exc
    return candidate


def manifest_source(entry: Any, label: str) -> Source:
    if not isinstance(entry, dict):
        raise RuntimeError(f"{label} must be an object.")
    key = entry.get("key")
    title = entry.get("title")
    if not isinstance(key, str) or not re.fullmatch(r"[a-z0-9][a-z0-9-]*", key):
        raise RuntimeError(f"{label} has an invalid key: {key!r}")
    if not isinstance(title, str) or not title.strip():
        raise RuntimeError(f"{label} has an invalid title.")
    # Source type is decided per entry. A Git source is additive: it does not
    # change how any filesystem entry resolves, and its repository is never
    # reused as a root for another document.
    source_type = entry.get("sourceType", FILESYSTEM_SOURCE_TYPE)
    if source_type == GIT_SOURCE_TYPE:
        spec = parse_git_source(entry, label)
        generated = write_generated_git_source(
            spec, GENERATED_DIR, key, title, label=f"{label} ({key})"
        )
        return Source(key, title, generated, spec)
    if source_type not in FILESYSTEM_SOURCE_TYPES:
        raise RuntimeError(f"{label} has an unknown sourceType: {source_type!r}")

    # A filesystem entry resolves exactly as it always has, under the fixed
    # redesign root. Git fields are rejected here so no filesystem document can
    # inherit another entry's repository or ref.
    borrowed = sorted({"repository", "ref"} & set(entry))
    if borrowed:
        raise RuntimeError(
            f"{label} is a filesystem source but declares Git fields: "
            + ", ".join(borrowed)
        )
    source_path = resolve_manifest_path(entry.get("root"), entry.get("path"), label)
    return Source(key, title, source_path)


def dynamic_source_definitions() -> dict[str, tuple[str, str]]:
    dynamic: dict[str, tuple[str, str]] = {}
    for index, entry in enumerate(load_source_manifest()["dynamicDocuments"]):
        label = f"dynamicDocuments[{index}]"
        if not isinstance(entry, dict):
            raise RuntimeError(f"{label} must be an object.")
        key = entry.get("key")
        title = entry.get("title")
        if not isinstance(key, str) or not re.fullmatch(r"[a-z0-9][a-z0-9-]*", key):
            raise RuntimeError(f"{label} has an invalid key: {key!r}")
        if not isinstance(title, str) or not title.strip():
            raise RuntimeError(f"{label} has an invalid title.")
        if key in dynamic:
            raise RuntimeError(f"Duplicate dynamic document key: {key}")
        dynamic[key] = (key, title)
    required = {"latest-handoff", "tm-project-master-context"}
    if set(dynamic) != required:
        raise RuntimeError(
            "dynamicDocuments must contain exactly latest-handoff and "
            "tm-project-master-context."
        )
    return dynamic


def fixed_sources() -> list[Source]:
    manifest = load_source_manifest()
    sources = [
        manifest_source(entry, f"documents[{index}]")
        for index, entry in enumerate(manifest["documents"])
    ]
    phases = manifest["phaseDocuments"]
    first = phases.get("first")
    last = phases.get("last")
    if isinstance(first, bool) or not isinstance(first, int):
        raise RuntimeError("phaseDocuments.first must be an integer.")
    if isinstance(last, bool) or not isinstance(last, int) or first > last:
        raise RuntimeError("phaseDocuments.last must be an integer not less than first.")
    phase_dir = resolve_manifest_path(
        phases.get("root"), phases.get("directory"), "phaseDocuments"
    )
    if not phase_dir.is_dir():
        raise RuntimeError(f"Phase directory is missing: {phase_dir}")
    for number in range(first, last + 1):
        matches = sorted(phase_dir.glob(f"{number:02d}-*.md"))
        if len(matches) != 1:
            raise RuntimeError(
                f"Expected exactly one phase {number:02d} Markdown file; found {len(matches)}."
            )
        sources.append(Source(f"phase-{number:02d}", matches[0].stem, matches[0]))
    return sources


def expected_source_count() -> int:
    return len(fixed_sources()) + len(dynamic_source_definitions())


def discover_sources() -> list[Source]:
    master_context_path = build_master_context_source()
    sources = fixed_sources()
    dynamic = dynamic_source_definitions()
    handoff_dir = ROOT / "docs/agent-handoffs"
    handoffs = sorted(
        handoff_dir.glob("*.md"), key=lambda p: p.stat().st_mtime_ns, reverse=True
    )
    if not handoffs:
        raise RuntimeError(f"No Markdown handoff found in {handoff_dir}")
    latest_key, latest_title = dynamic["latest-handoff"]
    master_key, master_title = dynamic["tm-project-master-context"]
    sources.append(Source(latest_key, latest_title, handoffs[0]))
    sources.append(
        Source(master_key, master_title, master_context_path)
    )

    missing = [str(source.path) for source in sources if not source.path.is_file()]
    if missing:
        raise RuntimeError("Missing required source files:\n" + "\n".join(missing))
    expected = expected_source_count()
    if len(sources) != expected:
        raise RuntimeError(
            f"Expected {expected} sources from the manifest; discovered {len(sources)}"
        )
    if len({source.title.casefold() for source in sources}) != len(sources):
        raise RuntimeError("Duplicate Google Doc titles were discovered.")
    if len({source.key for source in sources}) != len(sources):
        raise RuntimeError("Duplicate planning-pack source keys were discovered.")
    verify_git_sources(sources)
    return sources


def current_source_snapshot(sources: list[Source], state: dict[str, Any]) -> dict[str, Any]:
    drive_ids = {
        key: record.get("id")
        for key, record in state.get("documents", {}).items()
        if isinstance(record, dict) and record.get("id")
    }
    return build_snapshot(sources, drive_ids)


def gate_source_isolation(
    sources: list[Source], state: dict[str, Any], snapshot_path: Path | None
) -> dict[str, Any]:
    """Log the resolved source table, then fail closed before touching Drive."""

    snapshot = current_source_snapshot(sources, state)
    logging.info("Resolved planning-pack sources:\n%s", render_table(snapshot))

    def is_within(root: Path, ancestor: Path) -> bool:
        return root == ancestor or ancestor in root.parents

    redesign_rooted = []
    for document in snapshot["documents"]:
        if document["source_type"] != "filesystem":
            continue
        root = Path(document["root"])
        if is_within(root, ROOT):
            redesign_rooted.append(document["key"])
            continue
        if is_within(root, GENERATED_DIR):
            # The two dynamic documents the updater generates itself.
            continue
        raise SourceIsolationError(
            f"{document['key']} resolves from {document['root']}, which is neither the "
            "redesign checkout nor the updater's generated directory. Refusing to publish."
        )

    logging.info(
        "%d documents resolve from the redesign checkout; %d from Git.",
        len(redesign_rooted),
        sum(1 for document in snapshot["documents"] if document["source_type"] == "git"),
    )

    if snapshot_path is not None:
        assert_source_isolation(
            snapshot_path, snapshot, GIT_SOURCED_KEYS, REQUIRED_GIT_SOURCES, ROTATING_KEYS
        )
        logging.info("Source isolation verified against %s.", snapshot_path)
    return snapshot


def verify_git_sources(sources: list[Source]) -> int:
    """Re-resolve every Git-sourced document and prove the artifact matches.

    Called after discovery and again immediately before publication, so a stale
    or hand-edited generated artifact stops the run rather than reaching Drive.
    """

    verified = 0
    for source in sources:
        if source.git_spec is None:
            continue
        provenance = verify_generated_document(
            source.path, source.git_spec, label=f"Git source ({source.key})"
        )
        logging.info(
            "Git source verified: %s <- %s:%s @ %s",
            source.title,
            provenance["ref"],
            provenance["path"],
            provenance["path_commit"][:9],
        )
        verified += 1
    return verified


def configure_logging(scheduled: bool) -> Path:
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    log_path = LOG_DIR / "latest.log"
    handlers: list[logging.Handler] = [logging.FileHandler(log_path, mode="w", encoding="utf-8")]
    if not scheduled:
        handlers.append(logging.StreamHandler(sys.stdout))
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=handlers,
        force=True,
    )
    return log_path


def process_is_running(pid: int) -> bool:
    if pid <= 0:
        return False
    if os.name == "nt":
        process_query_limited_information = 0x1000
        still_active = 259
        kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        kernel32.OpenProcess.argtypes = [
            ctypes.c_ulong,
            ctypes.c_bool,
            ctypes.c_ulong,
        ]
        kernel32.OpenProcess.restype = ctypes.c_void_p
        kernel32.GetExitCodeProcess.argtypes = [
            ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_ulong),
        ]
        kernel32.GetExitCodeProcess.restype = ctypes.c_bool
        kernel32.CloseHandle.argtypes = [ctypes.c_void_p]
        kernel32.CloseHandle.restype = ctypes.c_bool

        handle = kernel32.OpenProcess(
            process_query_limited_information,
            False,
            pid,
        )
        if not handle:
            # Access denied still proves that a process owns this PID.
            return ctypes.get_last_error() == 5
        try:
            exit_code = ctypes.c_ulong()
            if not kernel32.GetExitCodeProcess(
                handle,
                ctypes.byref(exit_code),
            ):
                return False
            return exit_code.value == still_active
        finally:
            kernel32.CloseHandle(handle)

    try:
        os.kill(pid, 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    return True


def read_lock_pid() -> int | None:
    try:
        payload = json.loads(LOCK_FILE.read_text(encoding="utf-8"))
        pid = int(payload.get("pid"))
        return pid if pid > 0 else None
    except (OSError, TypeError, ValueError, json.JSONDecodeError):
        return None


def acquire_lock() -> None:
    if LOCK_FILE.exists():
        age = max(0.0, time.time() - LOCK_FILE.stat().st_mtime)
        owner_pid = read_lock_pid()
        if owner_pid is not None and process_is_running(owner_pid):
            raise UpdateAlreadyRunning(
                "TM Planning Pack is already updating in the background. "
                "No second update was started."
            )
        if owner_pid is None and age < 30:
            raise UpdateAlreadyRunning(
                "TM Planning Pack has just started updating in the background. "
                "No second update was started."
            )
        LOCK_FILE.unlink(missing_ok=True)
    try:
        fd = os.open(LOCK_FILE, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
    except FileExistsError as exc:
        raise UpdateAlreadyRunning(
            "TM Planning Pack is already updating in the background. "
            "No second update was started."
        ) from exc
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as handle:
            handle.write(json.dumps({"pid": os.getpid(), "started": time.time()}))
    except Exception:
        LOCK_FILE.unlink(missing_ok=True)
        raise


def release_lock() -> None:
    if LOCK_FILE.exists() and read_lock_pid() == os.getpid():
        LOCK_FILE.unlink(missing_ok=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_style_font(style: Any, name: str, size: int, color: str, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def create_reference_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Arial", 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (20, "000000", 20, 6),
        "Heading 2": (16, "000000", 18, 6),
        "Heading 3": (14, "434343", 16, 4),
        "Heading 4": (12, "000000", 14, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        set_style_font(style, "Arial", size, color, bold=False)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, "Arial", 11, "000000")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    doc.add_paragraph("Reference document for Google Docs imports.")
    doc.save(REFERENCE_DOCX)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def normalize_docx(path: Path) -> None:
    doc = Document(path)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    style_specs = {
        "Normal": (11, "000000", 0, 8, False),
        "Heading 1": (20, "000000", 20, 6, False),
        "Heading 2": (16, "000000", 18, 6, False),
        "Heading 3": (14, "434343", 16, 4, False),
        "Heading 4": (12, "000000", 14, 4, False),
    }
    for style_name, (size, color, before, after, bold) in style_specs.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        set_style_font(style, "Arial", size, color, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if style_name == "Normal":
            style.paragraph_format.line_spacing = 1.15
        else:
            style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        set_style_font(style, "Arial", 11, "000000")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "9360")
        tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), "0")
        tbl_ind.set(qn("w:type"), "dxa")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margins(cell)

    doc.save(path)


def validate_docx(path: Path) -> None:
    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError(f"DOCX was not created: {path}")
    with zipfile.ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required - names
        if missing:
            raise RuntimeError(f"Invalid DOCX {path.name}; missing {sorted(missing)}")
    subprocess.run(
        [sys.executable, str(SANITIZER), str(path), "--check"],
        check=True,
        capture_output=True,
        text=True,
    )


def build_docx(source: Source, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    if source.path.suffix.casefold() == ".md":
        pypandoc.convert_file(
            str(source.path),
            "docx",
            outputfile=str(destination),
            extra_args=[
                f"--reference-doc={REFERENCE_DOCX}",
                f"--resource-path={source.path.parent}{os.pathsep}{ROOT}",
                "--wrap=preserve",
            ],
        )
    elif source.path.suffix.casefold() == ".docx":
        shutil.copy2(source.path, destination)
    else:
        raise RuntimeError(f"Unsupported source type: {source.path}")
    normalize_docx(destination)
    subprocess.run(
        [sys.executable, str(SANITIZER), str(destination), "--in-place"],
        check=True,
        capture_output=True,
        text=True,
    )
    validate_docx(destination)


def load_state() -> dict[str, Any]:
    if not STATE_FILE.exists():
        return {"version": 1, "folder_id": None, "folder_url": None, "documents": {}}
    with STATE_FILE.open("r", encoding="utf-8") as handle:
        state = json.load(handle)
    if state.get("version") != 1 or not isinstance(state.get("documents"), dict):
        raise RuntimeError(f"Unsupported or invalid state file: {STATE_FILE}")
    return state


def atomic_write_json(path: Path, data: dict[str, Any]) -> None:
    temp = path.with_suffix(path.suffix + ".tmp")
    with temp.open("w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2, sort_keys=True)
        handle.write("\n")
    os.replace(temp, path)


def get_credentials(scheduled: bool) -> Credentials:
    credentials: Credentials | None = None
    if TOKEN_FILE.exists():
        credentials = Credentials.from_authorized_user_file(str(TOKEN_FILE), SCOPES)
    if credentials and credentials.expired and credentials.refresh_token:
        credentials.refresh(Request())
    if not credentials or not credentials.valid:
        if scheduled:
            raise RuntimeError(
                "Google authorization is missing or expired. Double-click the Desktop updater once."
            )
        if not CREDENTIALS_FILE.is_file():
            raise RuntimeError(f"OAuth Desktop credential not found: {CREDENTIALS_FILE}")
        flow = InstalledAppFlow.from_client_secrets_file(str(CREDENTIALS_FILE), SCOPES)
        credentials = flow.run_local_server(port=0, open_browser=True, prompt="consent")
    TOKEN_FILE.write_text(credentials.to_json(), encoding="utf-8")
    try:
        os.chmod(TOKEN_FILE, 0o600)
    except OSError:
        pass
    return credentials


def q_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("'", "\\'")


def list_marked_files(service: Any, query: str) -> list[dict[str, Any]]:
    response = service.files().list(
        q=query,
        spaces="drive",
        fields="files(id,name,mimeType,parents,appProperties,webViewLink,trashed)",
        pageSize=100,
    ).execute()
    return response.get("files", [])


def get_or_create_folder(service: Any, state: dict[str, Any]) -> dict[str, Any]:
    folder_id = state.get("folder_id")
    if folder_id:
        try:
            folder = service.files().get(
                fileId=folder_id,
                fields="id,name,mimeType,webViewLink,trashed",
            ).execute()
            if not folder.get("trashed") and folder.get("mimeType") == GOOGLE_FOLDER_MIME:
                if folder.get("name") != FOLDER_NAME:
                    folder = service.files().update(
                        fileId=folder_id,
                        body={"name": FOLDER_NAME},
                        fields="id,name,mimeType,webViewLink,trashed",
                    ).execute()
                state["folder_url"] = folder.get("webViewLink") or f"https://drive.google.com/drive/folders/{folder_id}"
                atomic_write_json(STATE_FILE, state)
                return folder
        except HttpError as exc:
            if getattr(exc.resp, "status", None) not in (404, 410):
                raise

    marker = q_escape(APP_MARKER)
    query = (
        f"trashed = false and mimeType = '{GOOGLE_FOLDER_MIME}' "
        f"and appProperties has {{ key='tm_pack_owner' and value='{marker}' }}"
    )
    matches = list_marked_files(service, query)
    if len(matches) > 1:
        raise RuntimeError("Multiple app-managed planning-pack folders exist; refusing to guess.")
    if matches:
        folder = matches[0]
    else:
        folder = service.files().create(
            body={
                "name": FOLDER_NAME,
                "mimeType": GOOGLE_FOLDER_MIME,
                "appProperties": {"tm_pack_owner": APP_MARKER},
            },
            fields="id,name,mimeType,webViewLink,trashed",
        ).execute()
    state["folder_id"] = folder["id"]
    state["folder_url"] = folder.get("webViewLink") or f"https://drive.google.com/drive/folders/{folder['id']}"
    atomic_write_json(STATE_FILE, state)
    return folder


def recover_document(service: Any, folder_id: str, source: Source) -> dict[str, Any] | None:
    key = q_escape(source.key)
    query = (
        f"trashed = false and '{q_escape(folder_id)}' in parents "
        f"and appProperties has {{ key='tm_pack_key' and value='{key}' }}"
    )
    matches = list_marked_files(service, query)
    if len(matches) > 1:
        raise RuntimeError(f"Multiple app-managed Docs found for {source.key}; refusing to guess.")
    return matches[0] if matches else None


def mapped_document(service: Any, state: dict[str, Any], folder_id: str, source: Source) -> dict[str, Any] | None:
    record = state["documents"].get(source.key, {})
    file_id = record.get("id")
    if file_id:
        try:
            item = service.files().get(
                fileId=file_id,
                fields="id,name,mimeType,parents,appProperties,webViewLink,trashed,modifiedTime",
            ).execute()
            if (
                not item.get("trashed")
                and item.get("mimeType") == GOOGLE_DOC_MIME
                and folder_id in item.get("parents", [])
            ):
                return item
        except HttpError as exc:
            if getattr(exc.resp, "status", None) not in (404, 410):
                raise
    return recover_document(service, folder_id, source)


def upload_one(
    service: Any,
    state: dict[str, Any],
    folder_id: str,
    source: Source,
    docx_path: Path,
    source_hash: str,
) -> tuple[str, dict[str, Any]]:
    existing = mapped_document(service, state, folder_id, source)
    old_record = state["documents"].get(source.key, {})
    if existing and old_record.get("sha256") == source_hash and existing.get("name") == source.title:
        action = "unchanged"
        item = existing
    else:
        media = MediaFileUpload(str(docx_path), mimetype=DOCX_MIME, resumable=False)
        metadata = {
            "name": source.title,
            "appProperties": {
                "tm_pack_owner": APP_MARKER,
                "tm_pack_key": source.key,
            },
        }
        fields = "id,name,mimeType,parents,appProperties,webViewLink,trashed,modifiedTime"
        if existing:
            item = service.files().update(
                fileId=existing["id"],
                body=metadata,
                media_body=media,
                fields=fields,
            ).execute()
            action = "updated"
        else:
            metadata["mimeType"] = GOOGLE_DOC_MIME
            metadata["parents"] = [folder_id]
            item = service.files().create(
                body=metadata,
                media_body=media,
                fields=fields,
            ).execute()
            action = "created"

    if item.get("mimeType") != GOOGLE_DOC_MIME:
        raise RuntimeError(f"{source.title} is not a native Google Doc after upload.")
    if folder_id not in item.get("parents", []):
        raise RuntimeError(f"{source.title} is not in the managed Drive folder.")
    record = {
        "id": item["id"],
        "title": source.title,
        "source": str(source.path),
        "sha256": source_hash,
        "webViewLink": item.get("webViewLink") or f"https://docs.google.com/document/d/{item['id']}/edit",
        "modifiedTime": item.get("modifiedTime"),
    }
    state["documents"][source.key] = record
    atomic_write_json(STATE_FILE, state)
    return action, record


def prepare_sources(sources: list[Source], output_dir: Path) -> list[dict[str, str]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict[str, str]] = []
    for index, source in enumerate(sources, start=1):
        logging.info("Preparing %d/%d: %s", index, len(sources), source.title)
        destination = output_dir / f"{source.key}.docx"
        build_docx(source, destination)
        results.append(
            {
                "key": source.key,
                "title": source.title,
                "source": str(source.path),
                "source_sha256": sha256(source.path),
                "docx": str(destination),
                "docx_sha256": sha256(destination),
            }
        )
    return results


def run_update(scheduled: bool, source_snapshot: Path | None = None) -> dict[str, Any]:
    sources = discover_sources()
    expected_keys = {source.key for source in sources}
    state = load_state()
    gate_source_isolation(sources, state, source_snapshot)
    retired_keys = sorted(set(state["documents"]) - expected_keys)
    if retired_keys:
        raise RuntimeError(
            "The source catalog retired managed keys without an explicit Drive "
            "retirement operation: " + ", ".join(retired_keys)
        )
    if not REFERENCE_DOCX.exists():
        create_reference_docx()
    credentials = get_credentials(scheduled)
    service = build("drive", "v3", credentials=credentials, cache_discovery=False)
    about = service.about().get(fields="importFormats").execute()
    targets = about.get("importFormats", {}).get(DOCX_MIME, [])
    if GOOGLE_DOC_MIME not in targets:
        raise RuntimeError("This Google Drive account does not advertise DOCX-to-Google-Docs import support.")

    # Publication gate: nothing reaches Drive unless every Git-sourced document
    # still matches the ref it claims, re-checked here rather than trusted from
    # discovery time.
    verify_git_sources(sources)

    folder = get_or_create_folder(service, state)
    folder_id = folder["id"]
    counts = {"created": 0, "updated": 0, "unchanged": 0}
    documents: list[dict[str, Any]] = []
    with tempfile.TemporaryDirectory(prefix="tm_planning_pack_") as temp_dir:
        temp = Path(temp_dir)
        for index, source in enumerate(sources, start=1):
            source_hash = sha256(source.path)
            existing_record = state["documents"].get(source.key, {})
            existing = mapped_document(service, state, folder_id, source)
            if existing and existing_record.get("sha256") == source_hash and existing.get("name") == source.title:
                action, record = upload_one(
                    service, state, folder_id, source, temp / "unused.docx", source_hash
                )
            else:
                logging.info("Converting %d/%d: %s", index, len(sources), source.title)
                docx_path = temp / f"{source.key}.docx"
                build_docx(source, docx_path)
                action, record = upload_one(
                    service, state, folder_id, source, docx_path, source_hash
                )
            counts[action] += 1
            documents.append({"key": source.key, "action": action, **record})
            logging.info("%s: %s", action.capitalize(), source.title)

    stored_keys = set(state["documents"])
    if stored_keys != expected_keys:
        missing = sorted(expected_keys - stored_keys)
        raise RuntimeError(
            "Drive state does not contain the exact manifest source set. Missing: "
            + (", ".join(missing) if missing else "none")
        )
    summary = {
        "success": True,
        "completed_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
        "folder_id": folder_id,
        "folder_url": state["folder_url"],
        "document_count": len(documents),
        "counts": counts,
        "documents": documents,
    }
    atomic_write_json(SUMMARY_FILE, summary)
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description="Refresh native Google Docs for the TM planning pack.")
    parser.add_argument("--scheduled", action="store_true", help="Run unattended and do not open the folder.")
    parser.add_argument("--authorize-only", action="store_true", help="Complete Google sign-in without changing Drive files.")
    parser.add_argument("--prepare-only", action="store_true", help="Convert and validate locally without Google access.")
    parser.add_argument("--output-dir", type=Path, default=None, help="Destination used with --prepare-only.")
    parser.add_argument(
        "--source-manifest",
        type=Path,
        default=None,
        help=(
            "Read the source catalog from this manifest file instead of the default. "
            "Selects the manifest file only; it does not re-root any document."
        ),
    )
    parser.add_argument(
        "--source-snapshot",
        type=Path,
        default=None,
        help="Fail before writing to Drive unless every resolved source matches this snapshot.",
    )
    parser.add_argument(
        "--print-sources",
        action="store_true",
        help="Print the resolved source-resolution table and exit without Google access.",
    )
    args = parser.parse_args()
    try:
        acquire_lock()
    except UpdateAlreadyRunning as exc:
        if not args.scheduled:
            print(str(exc))
        return 0
    except Exception as exc:
        if not args.scheduled:
            print(f"TM Planning Pack update could not start: {exc}", file=sys.stderr)
        return 1

    log_path = configure_logging(args.scheduled)
    try:
        set_source_manifest(args.source_manifest)
        sources = discover_sources()
        if args.print_sources:
            snapshot = gate_source_isolation(sources, load_state(), args.source_snapshot)
            print(render_table(snapshot))
            print(json.dumps(snapshot, indent=2, sort_keys=True))
            return 0
        if not REFERENCE_DOCX.exists():
            create_reference_docx()
        if args.authorize_only:
            credentials = get_credentials(False)
            service = build("drive", "v3", credentials=credentials, cache_discovery=False)
            service.about().get(fields="importFormats").execute()
            logging.info("Google authorization and Drive API access verified. No Drive files were changed.")
            return 0
        if args.prepare_only:
            if args.output_dir is None:
                raise RuntimeError("--prepare-only requires --output-dir")
            results = prepare_sources(sources, args.output_dir.resolve())
            report = {
                "success": True,
                "mode": "prepare-only",
                "document_count": len(results),
                "documents": results,
            }
            atomic_write_json(args.output_dir.resolve() / "verification.json", report)
            logging.info("Prepared and structurally verified %d DOCX files.", len(results))
            return 0

        summary = run_update(args.scheduled, args.source_snapshot)
        logging.info(
            "Complete: %d created, %d updated, %d unchanged.",
            summary["counts"]["created"],
            summary["counts"]["updated"],
            summary["counts"]["unchanged"],
        )
        if not args.scheduled:
            folder_url = summary["folder_url"]
            try:
                if LOCAL_DRIVE_FOLDER.is_dir():
                    os.startfile(str(LOCAL_DRIVE_FOLDER))
                else:
                    webbrowser.open(folder_url)
            except OSError:
                webbrowser.open(folder_url)
        return 0
    except SourceIsolationError as exc:
        # A managed document resolved from somewhere other than its recorded
        # source. Publishing now would overwrite a stable Drive ID with the
        # wrong content, so stop before any Drive write.
        logging.error("Source isolation check failed: %s", exc)
        if not args.scheduled:
            print(f"TM Planning Pack update stopped: {exc}", file=sys.stderr)
        atomic_write_json(
            SUMMARY_FILE,
            {
                "success": False,
                "completed_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
                "error": str(exc),
                "error_kind": "source-isolation",
                "log": str(log_path),
            },
        )
        return 1
    except GitSourceError as exc:
        # A Git-sourced document could not be resolved or did not match its ref.
        # There is no filesystem fallback by design: stop rather than publish a
        # stale working-tree copy.
        logging.error("Git-sourced document failed to resolve: %s", exc)
        message = (
            f"TM Planning Pack update stopped: {exc}\n"
            "The document is sourced from Git and has no filesystem fallback. "
            "Fix the configured repository, ref, or path in "
            "docs/redesign/CLAUDE-PROJECT-SOURCES.json, or commit the missing "
            "content on the configured lineage, then run the updater again."
        )
        if not args.scheduled:
            print(message, file=sys.stderr)
        atomic_write_json(
            SUMMARY_FILE,
            {
                "success": False,
                "completed_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
                "error": str(exc),
                "error_kind": "git-source",
                "log": str(log_path),
            },
        )
        return 1
    except Exception as exc:
        logging.exception("Update failed: %s", exc)
        failure = {
            "success": False,
            "completed_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
            "error": str(exc),
            "log": str(log_path),
        }
        atomic_write_json(SUMMARY_FILE, failure)
        return 1
    finally:
        release_lock()


if __name__ == "__main__":
    raise SystemExit(main())
